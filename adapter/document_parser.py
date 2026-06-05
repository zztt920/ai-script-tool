"""多格式文档解析：从 .txt / .docx / .pdf 提取纯文本。"""

import io
import logging
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)


def parse_document(file_path: str | Path) -> str:
    """根据文件扩展名自动选择解析器并返回纯文本。

    Args:
        file_path: 文件路径

    Returns:
        解析后的纯文本内容

    Raises:
        ValueError: 不支持的格式或解析失败
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()
    try:
        if suffix == ".txt":
            return _parse_txt(path)
        elif suffix == ".docx":
            return _parse_docx(path)
        elif suffix == ".pdf":
            return _parse_pdf(path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}，仅接受 .txt / .docx / .pdf")
    except (ValueError, FileNotFoundError):
        raise
    except Exception as e:
        logger.exception("解析文件失败: %s", file_path)
        raise ValueError(f"文件解析失败 ({path.name}): {e}") from e


def validate_suffix(filename: str) -> bool:
    """验证文件名后缀是否在支持列表中。"""
    return Path(filename).suffix.lower() in {".txt", ".docx", ".pdf"}


def _parse_txt(path: Path) -> str:
    """读取纯文本文件，自动检测编码。"""
    # 按优先级尝试编码
    for enc in ["utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030", "latin-1"]:
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _parse_docx(path: Path) -> str:
    """从 .docx 提取完整文本（含段落和表格）。"""
    from docx import Document

    doc = Document(str(path))
    parts: List[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # 同时提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("  ".join(cells))

    if not parts:
        raise ValueError("Word 文档中未找到文本内容")
    return "\n".join(parts)


def _parse_pdf(path: Path) -> str:
    """从 PDF 提取文本，优先用 pdfplumber，回退到 pdfminer。"""
    try:
        import pdfplumber
        parts: List[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
        if not parts:
            raise ValueError("PDF 中未找到可提取的文本")
        return "\n\n".join(parts)
    except ImportError:
        pass

    # 回退方案: pdfminer.six
    from pdfminer.high_level import extract_text
    text = extract_text(str(path))
    if not text or not text.strip():
        raise ValueError("PDF 中未找到可提取的文本")
    return text
